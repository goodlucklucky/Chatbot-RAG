import os
import io
import json
import time
from datetime import datetime
import hashlib
import pytesseract
from PIL import Image
import pymupdf
import markdown2
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chat_models import init_chat_model
from langchain_core.messages import HumanMessage
from langchain_openai import OpenAIEmbeddings
from pinecone import Pinecone, ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from langgraph.checkpoint.memory import MemorySaver
from langgraph.prebuilt import create_react_agent
from langchain.tools.retriever import create_retriever_tool

memory = MemorySaver()
llm = init_chat_model("gpt-4o-mini", model_provider="openai")
embeddings = OpenAIEmbeddings(model="text-embedding-3-large")

# Pinecone implementation
pinecone_api_key = os.environ.get("PINECONE_API_KEY")

pc = Pinecone(api_key=pinecone_api_key)

def get_user_index(user_id: str):
    index_name = hashlib.md5(user_id.encode()).hexdigest()
    existing_indexes = [index_info["name"] for index_info in pc.list_indexes()]

    if index_name not in existing_indexes:
        pc.create_index(
            name=index_name,
            dimension=3072,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1"),
        )
        while not pc.describe_index(index_name).status["ready"]:
            time.sleep(1)

    return pc.Index(index_name)

def save_docx_file(content: str, user_id: str) -> str:
    filename = f"{user_id}_{int(time.time())}.docx"
    doc_path = os.path.join("downloads", filename)
    os.makedirs("downloads", exist_ok=True)

    # Convert markdown to HTML
    html = markdown2.markdown(content)
    soup = BeautifulSoup(html, "html.parser")

    doc = DocxDocument()

    for element in soup:
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == "p":
            para = doc.add_paragraph()
            # Check for inline formatting (bold, italic, code)
            for child in element.children:
                if isinstance(child, str):
                    para.add_run(child)
                elif child.name == "strong":  # Bold
                    para.add_run(child.text).bold = True
                elif child.name == "em":  # Italic
                    para.add_run(child.text).italic = True
                elif child.name == "code":  # Code
                    para.add_run(child.text).font.name = 'Courier New'
                    para.add_run(child.text).font.size = Pt(10)

        else:
            # For any other HTML element, just add the text
            doc.add_paragraph(element.text)

    doc.save(doc_path)
    current_doc_path = os.path.join("downloads", f"{user_id}_current.docx")
    doc.save(current_doc_path)

    # Replace with your actual base URL
    return f"{os.environ.get('NEXT_PUBLIC_BACKEND_URL', 'http://localhost:5000')}/downloads/{filename}"

def edit_docx(docx_path, changes):
    doc = DocxDocument(docx_path)

    for change in changes:
        para_idx = change["para_idx"]
        new_text = change["new_text"]
        para = doc.paragraphs[para_idx]
        para.text = new_text

    doc.save(docx_path)

def classify_intent(agent, prompt: str) -> str:
    instruction = (
        "You are an intent classifier for an AI-powered document editing assistant. "
        "Given the user's chat message and the current workflow, classify the user's intent as one of the following: "
        "[create_document, upload_document, generate_outline, select_section, suggest_section, confirm_section, "
        "edit_section, continue_editing, reject_change, accept_change, download_document, export_document, chat].\n"
        "Definitions:\n"
        "- create_document: User wants to start a new document from scratch.\n"
        "- upload_document: User wants to edit current uploading document.\n"
        "- generate_outline: User wants the AI to generate a draft/outline.\n"
        "- select_section: User wants to select or retrieve a section of the document.\n"
        "- suggest_section: AI suggests a section to the user for confirmation.\n"
        "- confirm_section: User confirms the suggested section is correct and user wants to modify this section in the future.\n"
        "- edit_section: User wants to edit the currently selected section.\n"
        "- continue_editing: User wants to make further changes to the current section.\n"
        "- reject_change: User wants to discard the last proposed change.\n"
        "- accept_change: User wants to accept and apply the last proposed change.\n"
        "- download_document: User wants to download the document.\n"
        "- export_document: User wants to export the document in another format.\n"
        "- chat: General conversation or unclear intent.\n"
        "Always respond with ONLY the intent word from the list above. "
        f"User message: \"{prompt}\""
    )
    config = {"configurable": {"thread_id": "intent_classifying"}}
    response = agent.invoke({"messages": [HumanMessage(content=instruction)]}, config=config)
    content = response.get("messages")[-1].content
    # Clean up and return only the intent word (lowercase, no punctuation)
    intent = content.strip().lower().replace(".", "")
    # Validate output
    valid_intents = {
        "create_document", "upload_document", "generate_outline", "select_section",
        "suggest_section", "confirm_section", "edit_section", "continue_editing",
        "reject_change", "accept_change", "download_document", "export_document", "chat"
    }
    if intent not in valid_intents:
        return "chat"
    return intent

def edit_docx(docx_path, para_idx, new_text):
    doc = DocxDocument(docx_path)
    para = doc.paragraphs[para_idx]
    para.text = new_text
    doc.save(docx_path)

def classify_intent(agent, prompt: str) -> str:
    instruction = (
        "You are an intent classifier for an AI-powered document editing assistant. "
        "Given the user's chat message and the current workflow, classify the user's intent as one of the following: "
        "[create_document, upload_document, generate_outline, select_section, suggest_section, confirm_section, "
        "edit_section, continue_editing, reject_change, accept_change, download_document, export_document, chat].\n"
        "Definitions:\n"
        "- create_document: User wants to start a new document from scratch.\n"
        "- upload_document: User wants to edit current uploading document.\n"
        "- generate_outline: User wants the AI to generate a draft/outline.\n"
        "- select_section: User wants to select or retrieve a section of the document.\n"
        "- suggest_section: AI suggests a section to the user for confirmation.\n"
        "- confirm_section: User confirms the suggested section is correct and user wants to modify this section in the future.\n"
        "- edit_section: User wants to edit the currently selected section.\n"
        "- continue_editing: User wants to make further changes to the current section.\n"
        "- reject_change: User wants to discard the last proposed change.\n"
        "- accept_change: User wants to accept and apply the last proposed change.\n"
        "- download_document: User wants to download the document.\n"
        "- export_document: User wants to export the document in another format.\n"
        "- chat: General conversation or unclear intent.\n"
        "Always respond with ONLY the intent word from the list above. "
        f"User message: \"{prompt}\""
    )
    config = {"configurable": {"thread_id": "intent_classifying"}}
    response = agent.invoke({"messages": [HumanMessage(content=instruction)]}, config=config)
    content = response.get("messages")[-1].content
    # Clean up and return only the intent word (lowercase, no punctuation)
    intent = content.strip().lower().replace(".", "")
    # Validate output
    valid_intents = {
        "create_document", "upload_document", "generate_outline", "select_section",
        "suggest_section", "confirm_section", "edit_section", "continue_editing",
        "reject_change", "accept_change", "download_document", "export_document", "chat"
    }
    if intent not in valid_intents:
        return "chat"
    return intent

def load_documents(file_path: str):
    if file_path.lower().endswith(".pdf"):
        docs = []

        # If there are images in the PDF, process them using OCR
        print("Extracting images from PDF and performing OCR...")
        doc = pymupdf.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # Extract text from the page
            text = page.get_text()
            docs.append(Document(page_content=text))

            # Extract images from the page
            img_list = page.get_images(full=True)
            for img_index, img in enumerate(img_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                # Convert image bytes to PIL Image for OCR
                image = Image.open(io.BytesIO(image_bytes))

                # Run OCR on the image
                text_from_image = pytesseract.image_to_string(image)

                # Optionally, append OCR-extracted text as a document
                docs.append(Document(page_content=text_from_image))

        return docs
    elif file_path.lower().endswith((".png", ".jpg", ".jpeg")):
        print("Using OCR to extract text from image")
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return [Document(page_content=text)]
    else:
        print("Unsupported file format")
        return []

pending_edits = {}

def get_docx_sections(docx_path):
    doc = DocxDocument(docx_path)
    return [para.text for para in doc.paragraphs]

def invoke_stream(question: str, user_id: str, file_path: str, start: str, end: str, content: str):
    if start and end and content:
        pending_edits[user_id] = {
            "start": start,
            "end": end,
            "content": content,
            "new_content": None,
        }

    if os.path.exists(file_path):
        print("Loading and chunking contents of the file")
        if file_path.endswith(".doc") or file_path.endswith(".docx"):
            docs = DocxDocument(file_path)
            docs = [Document(page_content=para.text) for para in docs.paragraphs]
        else:
            docs = load_documents(file_path)
    else:
        docs = []

    index = get_user_index(user_id)

    vector_store = PineconeVectorStore(index=index, embedding=embeddings)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=200)
    all_splits = text_splitter.split_documents(docs)

    # Update metadata (illustration purposes)
    total_documents = len(all_splits)
    third = total_documents // 3

    for i, document in enumerate(all_splits):
        if i < third:
            document.metadata["section"] = "beginning"
        elif i < 2 * third:
            document.metadata["section"] = "middle"
        else:
            document.metadata["section"] = "end"


    print("Indexing chunks")
    _ = vector_store.add_documents(all_splits)
    print("Indexing finished!")

    retriever = vector_store.as_retriever()

    ### Build retriever tool ###
    tool = create_retriever_tool(
        retriever,
        user_id,
        "Independent chat room"
    )
    tools = [tool]

    agent_executor = create_react_agent(llm, tools, checkpointer=memory)
    config = {"configurable": {"thread_id": user_id}}

    full_content = ""  # <-- Accumulate content here
    gen_guide = (
        "You're an AI assistant helping user write the document."
        "If the user wants to generate or edit a document file, "
        "include ONLY the formal document content between the --- and --- markers. "
        "Do NOT break the markers across multiple tokens or stream them character-by-character. "
        "Only wrap formal document content in these markers. Keep anything else outside."
        "Format the document with proper headings, bullet points, or numbered sections if needed. \n\n"
    )
    guide = (
        "You are an AI assistant for document editing. Your job is to help the user select, review, and edit sections of a document interactively. Follow these instructions for all interactions:\n\n"
        "1. **Section Selection**\n"
        "   - If the user asks to select or review a section, analyze the document structure and their request.\n"
        "   - Identify the most relevant section or paragraph.\n"
        "   - Clearly display the selected section's content to the user, including its index or identifier if possible.\n"

        "2. **Editing**\n"
        "   - If the user provides an edit instruction (e.g., \"make this more concise\", \"change the tone\", etc.), apply the requested changes ONLY to the previously selected or discussed section.\n"
        "   - Present the revised text clearly, and mark it as the new version for that section.\n"

        "3. **Confirmation**\n"
        "   - After suggesting an edit, always ask the user if they want to accept or reject the change.\n"
        "   - If the user accepts, confirm the change and state that it will be applied to the document.\n"
        "   - If the user rejects, revert to the previous version and ask for further instructions if needed.\n"

        "4. **Formatting**\n"
        "   - When returning edited document content, always wrap the formal document content between lines containing only three dashes (`---`), with nothing else on those lines. Do not break the marker across multiple tokens or stream it character by character.\n"
        "   - All other explanations, confirmations, or questions should be outside the `---` markers.\n"
        "   - Do not format the document.\n"

        "5. **State Management**\n"
        "   - Remember the current section being discussed and its latest version based on the conversation history.\n"
        "   - Do not ask the user to repeat information already provided in the chat history.\n"

        "6. **Examples:**\n"
        "   - User: \"Edit the introduction to be more formal.\""
        '- AI: "Here is the revised introduction:\n---\n[Revised introduction text]\n---\nWould you like to accept this change?"'

        '- User: "Yes, apply it."'
        '- AI: "✅ Change applied to the introduction. Let me know if you want to edit another section."'

        '- User: "No, revert."'
        '- AI: "❌ Change discarded. The introduction remains as before. What would you like to do next?"'

        "7. **General Rules:**\n"
        "- Always be clear about which section you are editing.\n"
        "- Always use the `---` markers for document content.\n"
        "- Never lose track of the current section or the user's last instruction.\n"
        "- If unsure, ask clarifying questions.\n"
    )
    doc_writing_flag = False
    edit_flag = False
    download_flag = False
    prompt = question
    intent = classify_intent(agent_executor, question)
    if intent in ["create_document", "download_document", "export_document"]:
        prompt = gen_guide + question
        download_flag = True
    elif start and end and content:
        prompt = guide + "\n\nHere's the content of the selected document: \n\n" + content + question

    if intent in ["edit_section", "continue_editing"]:
        edit_info = pending_edits.get(user_id)
        if not edit_info or "content" not in edit_info:
            yield "Please select a section you want to edit."
            return
        
        orig_text = edit_info["content"]
        prompt = (
            "You are an AI assistant for document editing. "
            "The user wants to edit a section or paragraph of a document."
            "First, explain briefly what you changed or how you improved the content."
            "Then, present ONLY the formal edited content between lines containing only three dashes (`---`)."
            "Do NOT break the markers across multiple tokens or stream them character by character."
            "Keep explanations, confirmations, or questions outside the `---` markers."
            "Do not format the document content.\n\n"
            f"Edit the following content according to the instruction: {question}\nContent: {orig_text}\n"
            "After presenting the revised content, ask the user if they want to accept or reject this change."
        )
        edit_flag = True

    # if not os.path.exists(docx_path) and intent in ["select_section", "suggest_section", "confirm_section", "edit_section", "continue_editing", "reject_change", "accept_change"]:
    #     yield "I need a document to work with. Please upload one."
    #     return

    # if intent == "accept_change":
    #     download_url = f"{os.environ.get('NEXT_PUBLIC_BACKEND_URL', 'http://localhost:5000')}/downloads/{user_id}_current.docx"
    #     download_link = f"\n📄 Download your DOCX({download_url})"
    #     yield f"✅ Change applied successfully! The document has been updated.{download_link}"
    #     return

    # if intent in ["select_section", "suggest_section"]:
    #     para_idx, para_text = select_section_logic(agent_executor, docx_path, user_id, question)
    #     if para_text:
    #         pending_edits[user_id] = {
    #             "para_idx": para_idx,
    #             "orig_text": para_text,
    #             "new_text": None,
    #             "updated_date": datetime.now().isoformat()
    #         }
    #         yield f"Do you want to modify this section?\r\n\r\n{para_text}\r\n\r\n."
    #     else:
    #         yield "Sorry, I couldn't identify the section. Please clarify your request."
    #     return
    # if intent == "confirm_section":
    #     edit_info = pending_edits.get(user_id)
    #     if edit_info and edit_info.get("orig_text"):
    #         edit_info["updated_date"] = datetime.now().isoformat()
    #         yield f"Now, please tell me about what you want to change.\r\n\r\n{edit_info['orig_text']}\r\n\r\n."
    #     else:
    #         yield "No section to confirm. Please select a section first."
    #     return
    # if intent in ["edit_section", "continue_editing"]:
    #     edit_info = pending_edits.get(user_id)
    #     if not edit_info or "para_idx" not in edit_info:
    #         yield "Please select a section you want to edit."
    #         return
        
    #     orig_text = edit_info["orig_text"]
    #     prompt = (
    #         "If the user wants to edit a section or paragraph text, "
    #         "include ONLY the formal edited content between the --- and --- markers. "
    #         "Do NOT break the markers across multiple tokens or stream them character-by-character. "
    #         "Only wrap formal edited content in these markers. Keep anything else outside."
    #         "Format the document content with proper headings, bullet points, or numbered sections if needed."
    #         f"Edit the following paragraph according to the instruction: {question}\r\nParagraph: {orig_text}"
    #     )
    #     edit_flag = True

    # # Accept Change Flow
    # if intent == "accept_change":
    #     edit_info = pending_edits.get(user_id)
    #     if not edit_info or "new_text" not in edit_info:
    #         yield "No pending changes to accept."
    #         return
        
        # # Apply the edit
        # edit_docx(docx_path, edit_info["para_idx"], edit_info["new_text"])
        # download_url = f"{os.environ.get('NEXT_PUBLIC_BACKEND_URL', 'http://localhost:5000')}/downloads/{user_id}_current.docx"
        # download_link = f"\n📄 Download your DOCX({download_url})"
        
    #     # Clear pending edits
    #     del pending_edits[user_id]
    #     yield f"✅ Change applied successfully! The document has been updated.{download_link}"

    # # Reject Change Flow 
    # if intent == "reject_change":
    #     if user_id in pending_edits:
    #         # Clear the proposed edit but keep section selection
    #         if "new_text" in pending_edits[user_id]:
    #             del pending_edits[user_id]["new_text"]
    #         yield "❌ Change discarded. The document remains unchanged."
    #     else:
    #         yield "No pending changes to reject."
    #     return

    for message_chunk, metadata in agent_executor.stream(
        {"messages": [HumanMessage(content=prompt)]}, config=config, stream_mode="messages"
    ):
        if message_chunk.content and metadata["langgraph_node"] == "agent":
            content = message_chunk.content
            if download_flag and "---" == content.strip() and doc_writing_flag:
                doc_writing_flag = False
                download_url = save_docx_file(full_content.strip(), user_id)
                download_link = f'\n📄 Download your DOCX({download_url})'
                full_content = ""
                yield download_link
            if "---" == content.strip() and edit_flag and doc_writing_flag:
                doc_writing_flag = False
                pending_edits[user_id]["new_content"] = full_content
                full_content = ""
            if doc_writing_flag:
                full_content += content
                if edit_flag:
                    yield content
            elif "---" != content.strip() or download_flag == False:
                yield content
            if (edit_flag or download_flag) and "---" == content.strip():
                doc_writing_flag = True

    if intent in ["edit_section", "continue_editing"]:
        backend_url = os.environ.get('NEXT_PUBLIC_BACKEND_URL', 'http://localhost:5000')
        content_url = f"{backend_url}/apply_change?user_id={user_id}"
        yield f"\nApply your change({content_url})"

def get_pending_edits(user_id: str):
    return pending_edits[user_id]
